"""
DOCX extractor for the Multilingual RAG Ingestion Pipeline.

Extracts text from Microsoft Word (.docx) files, preserving
headings, paragraphs, lists, and tables.
"""

from pathlib import Path
from typing import List, Union, Optional
import re
from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
DOCX_AVAILABLE = True
from .base import BaseExtractor, ExtractionError
from ..models.blocks import ExtractedBlock
from ..models.enums import DocumentFormat, BlockType
from ..pipeline.helpers.id_generator import generate_block_id

class DOCXExtractor(BaseExtractor):
    """
    Extractor for Microsoft Word (.docx) files.
    
    Uses python-docx to parse documents and extract:
    - Headings (with levels)
    - Paragraphs
    - Lists
    - Tables
    """
    
    SUPPORTED_EXTENSIONS = [".docx"]
    FORMAT = DocumentFormat.DOCX
    
    # Mapping of Word style names to heading levels
    HEADING_STYLES = {
        'Title': 1,
        'Heading 1': 1, 'Heading1': 1,
        'Heading 2': 2, 'Heading2': 2,
        'Heading 3': 3, 'Heading3': 3,
        'Heading 4': 4, 'Heading4': 4,
        'Heading 5': 5, 'Heading5': 5,
        'Heading 6': 6, 'Heading6': 6,
    }
    
    def __init__(self, min_text_length: int = 1):
        """
        Initialize DOCX extractor.
        
        Args:
            min_text_length: Minimum text length for a block.
        """
        super().__init__()        
        self.min_text_length = min_text_length
    
    def extract(self, file_path: Union[str, Path]) -> List[ExtractedBlock]:
        """
        Extract blocks from a DOCX file.
        
        Args:
            file_path: Path to the .docx file.
            
        Returns:
            List of ExtractedBlock objects.
        """
        path = self._validate_file(file_path)
        
        try:
            doc = DocxDocument(path)
            return self._parse_document(doc)
        except Exception as e:
            raise ExtractionError(
                f"Failed to extract DOCX from {path}: {str(e)}",
                file_path=str(path),
                cause=e
            )
    
    def _parse_document(self, doc: "DocxDocument") -> List[ExtractedBlock]:
        """
        Parse DOCX document into blocks.
        
        Args:
            doc: python-docx Document object.
            
        Returns:
            List of ExtractedBlock objects.
        """
        blocks = []
        offset = 0
        
        # Process document body elements in order
        for element in doc.element.body:
            # Handle paragraphs
            if element.tag.endswith('}p'):
                block = self._process_paragraph_element(element, doc, offset)
                if block:
                    blocks.append(block)
                    offset += len(block.raw_text)
            
            # Handle tables
            elif element.tag.endswith('}tbl'):
                block = self._process_table_element(element, doc, offset)
                if block:
                    blocks.append(block)
                    offset += len(block.raw_text)
        
        return blocks
    
    def _process_paragraph_element(
        self,
        element,
        doc: "DocxDocument",
        offset: int
    ) -> Optional[ExtractedBlock]:
        """
        Process a paragraph element.
        
        Args:
            element: XML element.
            doc: Document object.
            offset: Current offset.
            
        Returns:
            ExtractedBlock or None.
        """
        # Find corresponding paragraph object
        para = None
        for p in doc.paragraphs:
            if p._element is element:
                para = p
                break
        
        if para is None:
            return None
        
        text = para.text.strip()
        
        if len(text) < self.min_text_length:
            return None
        
        # Determine block type and heading level
        block_type = BlockType.PARAGRAPH
        heading_level = None
        
        style_name = para.style.name if para.style else None
        
        # Check for heading style
        if style_name:
            for heading_style, level in self.HEADING_STYLES.items():
                if heading_style.lower() in style_name.lower():
                    block_type = BlockType.HEADING
                    heading_level = level
                    break
        
        # Check for list
        if self._is_list_paragraph(para):
            block_type = BlockType.LIST
        
        metadata = {
            "source_format": "docx",
            "style": style_name,
        }
        
        return ExtractedBlock(
            block_id=generate_block_id(),
            raw_text=text,
            block_type=block_type,
            source_offset=offset,
            heading_level=heading_level,
            metadata=metadata,
        )
    
    def _process_table_element(
        self,
        element,
        doc: "DocxDocument",
        offset: int
    ) -> Optional[ExtractedBlock]:
        """
        Process a table element.
        
        Args:
            element: XML element.
            doc: Document object.
            offset: Current offset.
            
        Returns:
            ExtractedBlock or None.
        """
        # Find corresponding table object
        table = None
        for t in doc.tables:
            if t._element is element:
                table = t
                break
        
        if table is None:
            return None
        
        # Extract table content
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                cells.append(cell_text)
            rows.append(' | '.join(cells))
        
        text = '\n'.join(rows)
        
        if len(text) < self.min_text_length:
            return None
        
        return ExtractedBlock(
            block_id=generate_block_id(),
            raw_text=text,
            block_type=BlockType.TABLE,
            source_offset=offset,
            metadata={
                "source_format": "docx",
                "rows": len(table.rows),
                "columns": len(table.columns) if table.rows else 0,
            },
        )
    
    def _is_list_paragraph(self, para: "Paragraph") -> bool:
        """
        Check if paragraph is a list item.
        
        Args:
            para: Paragraph object.
            
        Returns:
            True if paragraph is a list item.
        """
        # Check for numPr (numbering properties) in paragraph XML
        numPr = para._element.find(qn('w:numPr'))
        if numPr is not None:
            return True
        
        # Check style name
        style_name = para.style.name.lower() if para.style else ''
        list_indicators = ['list', 'bullet', 'number']
        
        return any(indicator in style_name for indicator in list_indicators)